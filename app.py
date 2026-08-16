"""
Smart 1 Sales Builder — Proposals & Insertion Orders
==========================================
Combined dashboard + proposal builder backend.

- Stores every quote (full builder state + generated PDF) in a database
  (SQLite by default; set DATABASE_URL for Render Postgres).
- Assigns quote numbers in the Q-10200 series (own counter, IO-style).
- Generates the Smart 1–branded proposal PDF (reportlab) and a formatted
  Word (.docx) copy on demand; the latest PDF is archived with the quote.
- Computes "what's still missing before this can become an IO" gap lists.
- Optional AI routes (draft a proposal from one sentence, rewrite copy)
  using the same OpenAI Responses API pattern as the IO builder.
- The existing IO builder app is NOT touched: AI estimate/description/
  landing-review/media-mix calls and the IO conversion (order number,
  client + internal PDFs, Smart 1 Suite webhook) go straight to the IO
  API (IO_API_BASE), which stays exactly as deployed today.
"""

import base64
import json
import os
import re
import threading
import logging
import zipfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import requests
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from flask_cors import CORS
from sqlalchemy import (Column, DateTime, Integer, LargeBinary, String, Text,
                        create_engine, func, or_)
from sqlalchemy.orm import declarative_base, sessionmaker

# ---- PDF (reportlab, same stack as the IO app) ----
from reportlab.lib import colors as rl_colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)
from xml.sax.saxutils import escape as xml_escape

# ---- Image tools (Pillow ships with reportlab; Cloudinary is optional) ----
from PIL import Image, ImageOps

try:
    import cloudinary
    import cloudinary.api
    import cloudinary.uploader
    _CLOUDINARY_IMPORTED = True
except Exception:  # pragma: no cover - only if the package is missing
    _CLOUDINARY_IMPORTED = False

# ---- Word export ----
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

load_dotenv()

# =====================================================================
# Version — shown in the app's top-right corner so a screenshot alone
# tells you exactly which build a salesperson is looking at.
# Bump APP_VERSION on every deploy that changes behavior.
# =====================================================================
APP_VERSION = "1.1.0"
BUILD_DATE = "2026-08-16"
VERSION_NOTES = "Adds Tools menu: SEO image resizer, namer & alt-text writer"

BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__, template_folder=str(BASE_DIR / "templates"))
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

CORS(app, resources={r"/api/*": {"origins": "*"}}, methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"])

# Cloudinary reads CLOUDINARY_URL from the environment. Optional: without it
# the image tool still optimizes and returns a ZIP, it just can't host files.
if _CLOUDINARY_IMPORTED:
    try:
        cloudinary.config(secure=True)
    except Exception:
        logger.warning("Cloudinary configuration failed; image hosting disabled.")

# Uploads are held in memory only. Keep a sane ceiling so a huge drop can't
# exhaust the dyno (Render Starter = 512 MB).
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_UPLOAD_MB", "60")) * 1024 * 1024

# =====================================================================
# Database (SQLite locally, Postgres on Render via DATABASE_URL)
# =====================================================================
_db_url = os.getenv("DATABASE_URL", "").strip()
if _db_url.startswith("postgres://"):  # Render's URL form -> SQLAlchemy form
    _db_url = _db_url.replace("postgres://", "postgresql://", 1)
if not _db_url:
    _db_url = "sqlite:///" + str(BASE_DIR / "smart1_sales_builder.db")
engine = create_engine(_db_url, future=True,
                       connect_args={"check_same_thread": False} if _db_url.startswith("sqlite") else {})
SessionLocal = sessionmaker(bind=engine, expire_on_commit=False, future=True)
Base = declarative_base()

QUOTE_BASE = 10199  # first quote is Q-10200 (matches the IO number family)
VALID_STATUSES = ["Draft", "Sent", "Approved", "Lost", "Expired", "Converted"]


class Counter(Base):
    __tablename__ = "counters"
    name = Column(String(50), primary_key=True)
    value = Column(Integer, nullable=False, default=QUOTE_BASE)


class Quote(Base):
    __tablename__ = "quotes"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_number = Column(String(20), unique=True, nullable=False, index=True)
    status = Column(String(20), nullable=False, default="Draft", index=True)
    client = Column(String(200), default="")
    website = Column(String(300), default="")
    industry = Column(String(100), default="")
    salesperson = Column(String(120), default="")
    data = Column(Text, default="{}")           # full builder state (JSON)
    monthly_budget = Column(Integer, default=0)
    months = Column(Integer, default=1)
    total_budget = Column(Integer, default=0)
    package = Column(String(40), default="")
    products_summary = Column(String(500), default="")
    goals_summary = Column(String(300), default="")
    geo_summary = Column(String(300), default="")
    revision = Column(Integer, default=1)
    io_number = Column(String(20), default="")
    io_client_pdf_url = Column(String(600), default="")
    io_internal_pdf_url = Column(String(600), default="")
    converted_at = Column(DateTime, nullable=True)
    pdf_blob = Column(LargeBinary, nullable=True)
    pdf_filename = Column(String(300), default="")
    pdf_generated_at = Column(DateTime, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc),
                        onupdate=lambda: datetime.now(timezone.utc))


class Activity(Base):
    __tablename__ = "activity"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_id = Column(Integer, index=True, nullable=True)
    icon = Column(String(10), default="•")
    text = Column(String(500), default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


Base.metadata.create_all(engine)

_COUNTER_LOCK = threading.Lock()


def next_quote_number(db):
    """Allocate the next quote number: Q-10200, Q-10201, ... (own counter)."""
    with _COUNTER_LOCK:
        row = db.get(Counter, "quote_number")
        if row is None:
            row = Counter(name="quote_number", value=QUOTE_BASE)
            db.add(row)
        row.value = int(row.value) + 1
        db.flush()
        return f"Q-{row.value}"


def log_activity(db, quote_id, icon, text):
    db.add(Activity(quote_id=quote_id, icon=icon, text=text[:490]))


# =====================================================================
# Gap check — what a finished IO still needs that a proposal may not have
# (mirrors the required pieces of the IO flow; rule-based so it works
#  with or without AI)
# =====================================================================
def compute_gaps(state):
    gaps = []
    s = state or {}

    def blank(v):
        return v in (None, "", [], {})

    if blank(s.get("clientContactName")) and blank(s.get("clientContactEmail")):
        gaps.append({"key": "contact", "label": "Client contact name, email, and phone"})
    if blank(s.get("startDate")):
        gaps.append({"key": "dates", "label": "Exact campaign start date (and end date or ongoing)"})
    if blank(s.get("creativeSource")):
        gaps.append({"key": "creative", "label": "Creative assets — client provides, or Smart 1 builds (creative fee)"})
    tp = s.get("trackingPlan") or {}
    if blank(tp.get("primaryConversion")) or blank(tp.get("ga4")):
        gaps.append({"key": "tracking", "label": "Tracking plan (primary conversion, GA4, call tracking, verifier)"})
    if blank(s.get("landingUrl")):
        gaps.append({"key": "landing", "label": "Final landing page URL"})
    if blank(s.get("salesContact")) or blank(s.get("salesEmail")):
        gaps.append({"key": "sales", "label": "Smart 1 sales contact name and email"})
    if blank(s.get("managementFeeAck")):
        gaps.append({"key": "fees", "label": "Management fee confirmation (and creative fee if Smart 1 builds)"})
    return gaps


# Guardrails (subset of the IO builder's rules, checked server-side too)
def compute_guardrails(state):
    warns = []
    s = state or {}
    items = s.get("items") or []
    budget = float(s.get("budget") or 0)
    months = int(s.get("months") or 1)
    for it in items:
        cat = (it.get("category") or "").upper()
        dollars = float(it.get("dollars") or 0)
        if "SEARCH ENGINE MARKETING" in cat and 0 < dollars < 1500:
            warns.append(f"Search budget ${dollars:,.0f}/mo is below the $1,500 monthly minimum.")
    if budget > 0 and len(items) > 4 and budget / max(len(items), 1) < 750:
        warns.append("Budget is split across many products — consider fewer products for impact.")
    if months < 3:
        warns.append("Campaign term under 3 months — most Smart 1 programs need 3+ months to optimize.")
    if (s.get("creativeSource") or "").lower().startswith("smart 1") and not s.get("creativeFee"):
        warns.append("Smart 1 is building creative but no creative fee is included.")
    return warns


# =====================================================================
# Serialization
# =====================================================================
def quote_json(q, include_data=False):
    state = {}
    try:
        state = json.loads(q.data or "{}")
    except Exception:
        state = {}
    out = {
        "id": q.id,
        "quote_number": q.quote_number,
        "status": q.status,
        "client": q.client,
        "website": q.website,
        "industry": q.industry,
        "salesperson": q.salesperson,
        "monthly_budget": q.monthly_budget,
        "months": q.months,
        "total_budget": q.total_budget,
        "package": q.package,
        "products_summary": q.products_summary,
        "goals_summary": q.goals_summary,
        "geo_summary": q.geo_summary,
        "revision": q.revision,
        "io_number": q.io_number,
        "io_client_pdf_url": q.io_client_pdf_url,
        "io_internal_pdf_url": q.io_internal_pdf_url,
        "has_pdf": bool(q.pdf_blob),
        "pdf_filename": q.pdf_filename,
        "created_at": q.created_at.isoformat() if q.created_at else "",
        "updated_at": q.updated_at.isoformat() if q.updated_at else "",
        "converted_at": q.converted_at.isoformat() if q.converted_at else "",
        "gaps": compute_gaps(state),
        "guardrails": compute_guardrails(state),
    }
    if include_data:
        out["data"] = state
    return out


def summarize_into(q, state):
    """Pull list/summary columns out of the saved state."""
    q.client = str(state.get("client") or "")[:200]
    q.website = str(state.get("url") or "")[:300]
    q.industry = str(state.get("industry") or "")[:100]
    q.salesperson = str(state.get("salesContact") or "")[:120]
    q.months = int(state.get("months") or 1)
    sel = state.get("selectedPackage") or {}
    q.monthly_budget = int(sel.get("monthly") or state.get("budget") or 0)
    q.total_budget = int(sel.get("total") or (q.monthly_budget * q.months))
    q.package = str(sel.get("name") or "")[:40]
    items = state.get("items") or []
    q.products_summary = " · ".join([str(i.get("product") or "") for i in items])[:500]
    q.goals_summary = ", ".join(state.get("objectives") or [])[:300]
    geo = state.get("geo") or state.get("geoType") or ""
    if state.get("geoType") == "City/ZIP + Radius" and state.get("radius"):
        geo = f"{state.get('geo','')} + {state.get('radius')} mi"
    q.geo_summary = str(geo)[:300]


# =====================================================================
# Routes — core
# =====================================================================
def _cloudinary_ready():
    """True when CLOUDINARY_URL (or the discrete vars) are configured."""
    if not _CLOUDINARY_IMPORTED:
        return False
    try:
        cfg = cloudinary.config()
        return bool(cfg.cloud_name and cfg.api_key and cfg.api_secret)
    except Exception:
        return False


@app.get("/health")
def health():
    return jsonify({
        "status": "ok",
        "version": APP_VERSION,
        "build_date": BUILD_DATE,
        "database": _db_url.split("://")[0],
        "openai_configured": bool(os.getenv("OPENAI_API_KEY")),
        "cloudinary_configured": _cloudinary_ready(),
        "io_api_base": os.getenv("IO_API_BASE", "https://insertionordersmart.onrender.com"),
    })


@app.get("/")
def index():
    return render_template("index.html", app_version=APP_VERSION)


@app.get("/api/config")
def api_config():
    return jsonify({
        "version": APP_VERSION,
        "build_date": BUILD_DATE,
        "version_notes": VERSION_NOTES,
        "database": _db_url.split("://")[0],
        "io_api_base": os.getenv("IO_API_BASE", "https://insertionordersmart.onrender.com"),
        "io_app_url": os.getenv("IO_APP_URL", "https://insertionordersmart.onrender.com"),
        "ai_enabled": bool(os.getenv("OPENAI_API_KEY")),
        "cloudinary_enabled": _cloudinary_ready(),
    })


# ---- Quotes CRUD ----
@app.post("/api/quotes")
def create_quote():
    body = request.get_json(force=True) or {}
    state = body.get("data") or {}
    db = SessionLocal()
    try:
        q = Quote(quote_number=next_quote_number(db),
                  status=body.get("status") or "Draft",
                  data=json.dumps(state, ensure_ascii=False))
        summarize_into(q, state)
        db.add(q)
        db.flush()
        log_activity(db, q.id, "🆕", f"{q.quote_number} created — {q.client or 'new quote'}")
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.get("/api/quotes")
def list_quotes():
    qstr = (request.args.get("q") or "").strip().lower()
    status = (request.args.get("status") or "").strip()
    limit = min(int(request.args.get("limit") or 200), 500)
    db = SessionLocal()
    try:
        query = db.query(Quote)
        if status and status != "all":
            query = query.filter(Quote.status == status)
        if qstr:
            like = f"%{qstr}%"
            query = query.filter(or_(
                func.lower(Quote.quote_number).like(like),
                func.lower(Quote.client).like(like),
                func.lower(Quote.website).like(like),
                func.lower(Quote.products_summary).like(like),
                func.lower(Quote.io_number).like(like),
            ))
        rows = query.order_by(Quote.updated_at.desc()).limit(limit).all()
        return jsonify({"ok": True, "quotes": [quote_json(r) for r in rows]})
    finally:
        db.close()


@app.get("/api/quotes/<int:qid>")
def get_quote(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.put("/api/quotes/<int:qid>")
def update_quote(qid):
    body = request.get_json(force=True) or {}
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        if "data" in body:
            state = body.get("data") or {}
            q.data = json.dumps(state, ensure_ascii=False)
            summarize_into(q, state)
            q.revision = (q.revision or 1) + (1 if body.get("bump_revision") else 0)
        if "status" in body:
            new_status = body["status"]
            if new_status not in VALID_STATUSES:
                return jsonify({"ok": False, "error": "Invalid status"}), 400
            if new_status != q.status:
                log_activity(db, q.id, {"Sent": "📤", "Approved": "✅", "Lost": "❌",
                                        "Converted": "🔁", "Expired": "⏰"}.get(new_status, "•"),
                             f"{q.quote_number} → {new_status} — {q.client}")
            q.status = new_status
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/duplicate")
def duplicate_quote(qid):
    db = SessionLocal()
    try:
        src = db.get(Quote, qid)
        if not src:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(src.data or "{}")
        # New quotes start clean of decision/IO fields
        for k in ("startDate", "ioPayload",):
            state.pop(k, None)
        q = Quote(quote_number=next_quote_number(db), status="Draft",
                  data=json.dumps(state, ensure_ascii=False))
        summarize_into(q, state)
        db.add(q)
        db.flush()
        log_activity(db, q.id, "📋", f"{q.quote_number} duplicated from {src.quote_number} — {q.client}")
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/converted")
def mark_converted(qid):
    """Called by the convert wizard after the IO API has issued the order
    number and generated the IO PDFs. Links everything back to the quote."""
    body = request.get_json(force=True) or {}
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        q.io_number = str(body.get("io_number") or "")[:20]
        q.io_client_pdf_url = str(body.get("client_pdf_url") or "")[:600]
        q.io_internal_pdf_url = str(body.get("internal_pdf_url") or "")[:600]
        q.status = "Converted"
        q.converted_at = datetime.now(timezone.utc)
        if "data" in body:
            q.data = json.dumps(body["data"], ensure_ascii=False)
            summarize_into(q, body["data"])
        log_activity(db, q.id, "🔁",
                     f"{q.quote_number} → IO #{q.io_number} — {q.client}"
                     + (", sent to Smart 1 Suite" if body.get("submitted") else ""))
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q)})
    finally:
        db.close()


# ---- Dashboard ----
@app.get("/api/dashboard")
def dashboard():
    db = SessionLocal()
    try:
        rows = db.query(Quote).all()
        now = datetime.now(timezone.utc)

        def aware(dt):
            if dt is None:
                return now
            return dt if dt.tzinfo else dt.replace(tzinfo=timezone.utc)

        open_rows = [r for r in rows if r.status in ("Draft", "Sent")]
        sent_rows = [r for r in rows if r.status == "Sent"]
        approved = [r for r in rows if r.status == "Approved"]
        conv_month = [r for r in rows if r.status == "Converted" and r.converted_at
                      and aware(r.converted_at).month == now.month and aware(r.converted_at).year == now.year]
        decided_90 = [r for r in rows if r.status in ("Approved", "Converted", "Lost")
                      and (now - aware(r.updated_at)).days <= 90]
        won_90 = [r for r in decided_90 if r.status in ("Approved", "Converted")]
        avg_days_out = round(sum((now - aware(r.updated_at)).days for r in sent_rows) / len(sent_rows)) if sent_rows else 0

        acts = db.query(Activity).order_by(Activity.created_at.desc()).limit(12).all()
        stale = [r for r in sent_rows if (now - aware(r.updated_at)).days >= 7]
        nudges = []
        for r in stale[:3]:
            nudges.append({"tag": "FOLLOW-UP NUDGE", "text": f"{r.client} ({r.quote_number}) was sent "
                          f"{(now - aware(r.updated_at)).days} days ago with no response. Consider a follow-up."})
        for r in approved[:3]:
            g = compute_gaps(json.loads(r.data or "{}"))
            nudges.append({"tag": "READY TO CONVERT", "text": f"{r.client} ({r.quote_number}) is approved. "
                          + (f"{len(g)} answers are still needed before the IO is complete." if g
                             else "Everything needed for the IO is already on file.")})
        for r in [x for x in rows if x.status == "Draft"][:5]:
            for w in compute_guardrails(json.loads(r.data or "{}"))[:1]:
                nudges.append({"tag": "PRICING GUARDRAIL", "text": f"{r.client or r.quote_number}: {w}"})

        return jsonify({"ok": True, "stats": {
            "open_count": len(open_rows),
            "open_monthly": sum(r.monthly_budget or 0 for r in open_rows),
            "awaiting_count": len(sent_rows),
            "awaiting_avg_days": avg_days_out,
            "approved_count": len(approved),
            "converted_month_count": len(conv_month),
            "converted_month_monthly": sum(r.monthly_budget or 0 for r in conv_month),
            "win_rate_90d": round(100 * len(won_90) / len(decided_90)) if decided_90 else 0,
        }, "activity": [{"icon": a.icon, "text": a.text,
                         "when": a.created_at.isoformat() if a.created_at else ""} for a in acts],
            "nudges": nudges[:6]})
    finally:
        db.close()


# =====================================================================
# Branded proposal PDF
# =====================================================================
NAVY = rl_colors.HexColor("#14284b")
BLUE = rl_colors.HexColor("#1f63ae")
GOLD = rl_colors.HexColor("#e5a323")
LINE = rl_colors.HexColor("#d5dee9")
SOFT = rl_colors.HexColor("#eef3f8")
MUTED = rl_colors.HexColor("#53657a")


def _p(text, style):
    return Paragraph(xml_escape(str(text or "")).replace("\n", "<br/>"), style)


def _money(n):
    try:
        return "$" + f"{float(n):,.0f}"
    except Exception:
        return str(n)


def build_proposal_pdf(q, state):
    title = f"S1M Proposal - {q.quote_number} - {q.client or 'Client'}"
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch,
                            topMargin=0.55 * inch, bottomMargin=0.6 * inch, title=title)
    ss = getSampleStyleSheet()
    st_title = ParagraphStyle("T", parent=ss["Title"], textColor=NAVY, fontSize=21, leading=25, alignment=TA_CENTER, spaceAfter=2)
    st_sub = ParagraphStyle("Sub", parent=ss["BodyText"], textColor=MUTED, fontSize=10.5, alignment=TA_CENTER, spaceAfter=10)
    st_h2 = ParagraphStyle("H2", parent=ss["Heading2"], textColor=NAVY, fontSize=13, leading=16, spaceBefore=13, spaceAfter=5)
    st_body = ParagraphStyle("B", parent=ss["BodyText"], fontSize=9.5, leading=13, spaceAfter=5)
    st_small = ParagraphStyle("S", parent=ss["BodyText"], fontSize=8, leading=10.5, textColor=MUTED)

    story = []
    # Branded header band
    band = Table([[Paragraph('<font color="#ffffff"><b>SMART 1 MARKETING</b></font>', ParagraphStyle("bnd", fontSize=13, leading=16, alignment=TA_CENTER))]],
                 colWidths=[7.4 * inch], rowHeights=[0.42 * inch])
    band.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), NAVY), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story += [band, Spacer(1, 14),
              Paragraph(xml_escape(f"Marketing Proposal — {q.client or 'Client'}"), st_title),
              Paragraph(xml_escape(f"Quote {q.quote_number}  ·  {datetime.now().strftime('%B %d, %Y')}"
                                   + (f"  ·  Prepared by {state.get('salesContact')}" if state.get("salesContact") else "")), st_sub)]

    # Meta table
    sel = state.get("selectedPackage") or {}
    meta = [["Client", q.client or ""],
            ["Website", q.website or ""],
            ["Campaign Goals", ", ".join(state.get("objectives") or [])],
            ["Target Area", q.geo_summary or ""],
            ["Term", f"{q.months} months"],
            ["Monthly Investment", _money(sel.get("monthly") or q.monthly_budget)],
            ["Total Investment", _money(sel.get("total") or q.total_budget)]]
    t = Table(meta, colWidths=[1.55 * inch, 5.85 * inch])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, -1), SOFT), ("TEXTCOLOR", (0, 0), (0, -1), NAVY),
                           ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                           ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("VALIGN", (0, 0), (-1, -1), "TOP"),
                           ("LEFTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 5),
                           ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story += [t]

    # Sections (editable in the builder; sensible defaults if absent)
    for sec in state.get("sections") or []:
        if not sec.get("enabled", True):
            continue
        story.append(Paragraph(xml_escape(sec.get("title") or ""), st_h2))
        if sec.get("body"):
            story.append(_p(sec.get("body"), st_body))
        kind = sec.get("kind")
        if kind == "reach":
            est = state.get("estimates") or {}
            if est:
                rows = [["Est. Population", "Addressable Audience", "Households", "Devices"],
                        [f"{int(est.get('pop') or 0):,}", f"{int(est.get('aud') or 0):,}",
                         f"{int(est.get('hh') or 0):,}", f"{int(est.get('dev') or 0):,}"]]
                rt = Table(rows, colWidths=[1.85 * inch] * 4)
                rt.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                                        ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("GRID", (0, 0), (-1, -1), 0.4, LINE),
                                        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
                story.append(rt)
            auds = state.get("audiences") or []
            if auds:
                story.append(_p("Audience layers: " + ", ".join(auds), st_small))
        elif kind == "mediaplan":
            items = state.get("items") or []
            if items:
                rows = [["Product", "Category", "Rate", "Monthly", f"Total ({q.months} mo)"]]
                for i in items:
                    rows.append([_p(i.get("product") or "", st_small), _p(i.get("category") or "", st_small),
                                 _p((i.get("rate") or "Managed") if not i.get("rateValue")
                                    else f"{i.get('rate')} {i.get('rateValue')}", st_small),
                                 _money(i.get("dollars") or 0), _money((i.get("dollars") or 0) * q.months)])
                mt = Table(rows, colWidths=[2.3 * inch, 1.7 * inch, 1.1 * inch, 1.15 * inch, 1.15 * inch], repeatRows=1)
                mt.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8),
                                        ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                        ("ALIGN", (3, 1), (-1, -1), "RIGHT"), ("TOPPADDING", (0, 0), (-1, -1), 5),
                                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
                story.append(mt)
        elif kind == "packages":
            pkgs = state.get("packages") or []
            if pkgs:
                selname = (state.get("selectedPackage") or {}).get("name")
                head = ["", *[p.get("name", "") + (" ★" if p.get("name") == selname else "") for p in pkgs]]
                rows = [head,
                        ["Monthly", *[_money(p.get("monthly")) for p in pkgs]],
                        [f"Total ({q.months} mo)", *[_money(p.get("total")) for p in pkgs]],
                        ["Est. impressions / mo tier", *[f"{int(p.get('impr') or 0):,}" for p in pkgs]]]
                pt = Table(rows, colWidths=[1.9 * inch] + [1.83 * inch] * len(pkgs))
                style = [("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                         ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                         ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                         ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"), ("TEXTCOLOR", (0, 1), (0, -1), NAVY),
                         ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]
                for ci, p in enumerate(pkgs):
                    if p.get("name") == selname:
                        style.append(("BACKGROUND", (ci + 1, 1), (ci + 1, -1), rl_colors.HexColor("#fff5df")))
                pt.setStyle(TableStyle(style))
                story.append(pt)
                story.append(_p("★ recommended package", st_small))
        elif kind == "kpis":
            kpis = state.get("kpis") or []
            if kpis:
                story.append(_p("KPIs: " + ", ".join(kpis), st_body))

    # Footer
    story += [Spacer(1, 16),
              _p("Smart 1 Marketing  ·  smart1marketing.com  ·  This proposal is valid for 30 days. "
                 "Rates follow the current Smart 1 rate card; final schedules are confirmed on the insertion order.", st_small)]
    doc.build(story)
    return buf.getvalue(), title


@app.get("/api/quotes/<int:qid>/pdf")
def quote_pdf(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(q.data or "{}")
        ensure_sections(state)
        pdf_bytes, title = build_proposal_pdf(q, state)
        q.pdf_blob = pdf_bytes
        q.pdf_filename = title + ".pdf"
        q.pdf_generated_at = datetime.now(timezone.utc)
        log_activity(db, q.id, "📄", f"PDF generated — {title}.pdf")
        db.commit()
        return send_file(BytesIO(pdf_bytes), mimetype="application/pdf",
                         as_attachment=request.args.get("download") == "1",
                         download_name=title + ".pdf")
    finally:
        db.close()


@app.get("/api/quotes/<int:qid>/pdf/archived")
def quote_pdf_archived(qid):
    """The last saved copy from the database (what 'save the PDF to the DB' returns)."""
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q or not q.pdf_blob:
            return jsonify({"ok": False, "error": "No archived PDF for this quote yet"}), 404
        return send_file(BytesIO(q.pdf_blob), mimetype="application/pdf",
                         as_attachment=False, download_name=q.pdf_filename or "proposal.pdf")
    finally:
        db.close()


# =====================================================================
# Word (.docx) export
# =====================================================================
def build_proposal_docx(q, state):
    d = Document()
    for section in d.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    def hcolor(p, size, color, bold=True, center=False):
        run = p.runs[0] if p.runs else p.add_run("")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    NAVY_D = RGBColor(0x14, 0x28, 0x4B)
    MUTED_D = RGBColor(0x53, 0x65, 0x7A)

    p = d.add_paragraph("SMART 1 MARKETING")
    hcolor(p, 16, NAVY_D, center=True)
    p = d.add_paragraph(f"Marketing Proposal — {q.client or 'Client'}")
    hcolor(p, 20, NAVY_D, center=True)
    p = d.add_paragraph(f"Quote {q.quote_number}  ·  {datetime.now().strftime('%B %d, %Y')}")
    hcolor(p, 10, MUTED_D, bold=False, center=True)

    sel = state.get("selectedPackage") or {}
    table = d.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, val in [("Client", q.client), ("Website", q.website),
                       ("Campaign Goals", ", ".join(state.get("objectives") or [])),
                       ("Target Area", q.geo_summary), ("Term", f"{q.months} months"),
                       ("Monthly Investment", _money(sel.get("monthly") or q.monthly_budget)),
                       ("Total Investment", _money(sel.get("total") or q.total_budget))]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val or "")
        row[0].paragraphs[0].runs[0].font.bold = True

    for sec in state.get("sections") or []:
        if not sec.get("enabled", True):
            continue
        h = d.add_heading(sec.get("title") or "", level=2)
        for run in h.runs:
            run.font.color.rgb = NAVY_D
        if sec.get("body"):
            d.add_paragraph(str(sec.get("body")))
        kind = sec.get("kind")
        if kind == "mediaplan" and state.get("items"):
            t2 = d.add_table(rows=1, cols=4)
            t2.style = "Light Grid Accent 1"
            hdr = t2.rows[0].cells
            for i, htxt in enumerate(["Product", "Category", "Monthly", f"Total ({q.months} mo)"]):
                hdr[i].text = htxt
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for it in state.get("items") or []:
                row = t2.add_row().cells
                row[0].text = str(it.get("product") or "")
                row[1].text = str(it.get("category") or "")
                row[2].text = _money(it.get("dollars") or 0)
                row[3].text = _money((it.get("dollars") or 0) * q.months)
        elif kind == "packages" and state.get("packages"):
            pkgs = state.get("packages") or []
            t3 = d.add_table(rows=1, cols=1 + len(pkgs))
            t3.style = "Light Grid Accent 1"
            hdr = t3.rows[0].cells
            hdr[0].text = ""
            for i, pkg in enumerate(pkgs):
                hdr[i + 1].text = pkg.get("name", "")
                hdr[i + 1].paragraphs[0].runs[0].font.bold = True
            for label, key in [("Monthly", "monthly"), ("Total", "total")]:
                row = t3.add_row().cells
                row[0].text = label
                for i, pkg in enumerate(pkgs):
                    row[i + 1].text = _money(pkg.get(key))
        elif kind == "kpis" and state.get("kpis"):
            d.add_paragraph("KPIs: " + ", ".join(state.get("kpis") or []))
        elif kind == "reach" and state.get("estimates"):
            est = state.get("estimates") or {}
            d.add_paragraph(f"Estimated population {int(est.get('pop') or 0):,} · addressable audience "
                            f"{int(est.get('aud') or 0):,} · households {int(est.get('hh') or 0):,} · "
                            f"devices {int(est.get('dev') or 0):,}")

    p = d.add_paragraph("Smart 1 Marketing · smart1marketing.com · This proposal is valid for 30 days.")
    hcolor(p, 8, MUTED_D, bold=False, center=True)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


@app.get("/api/quotes/<int:qid>/docx")
def quote_docx(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(q.data or "{}")
        ensure_sections(state)
        blob = build_proposal_docx(q, state)
        log_activity(db, q.id, "📝", f"Word copy exported — {q.quote_number}")
        db.commit()
        return send_file(BytesIO(blob),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True,
                         download_name=f"S1M Proposal - {q.quote_number} - {q.client or 'Client'}.docx")
    finally:
        db.close()


# =====================================================================
# Default proposal sections (used when the builder hasn't customized them)
# =====================================================================
def ensure_sections(state):
    if state.get("sections"):
        return state
    client = state.get("client") or "the client"
    goals = ", ".join(state.get("objectives") or []) or "the campaign goals"
    state["sections"] = [
        {"id": "about", "title": f"About {client}", "kind": "text", "enabled": True,
         "body": state.get("description") or ""},
        {"id": "strategy", "title": "Campaign Strategy", "kind": "text", "enabled": True,
         "body": f"This program is built around {goals.lower()}, combining the Smart 1 products below to "
                 f"reach the right audience in {state.get('geo') or 'the target area'} with the frequency "
                 f"needed to drive results."},
        {"id": "reach", "title": "Target Audience & Reach", "kind": "reach", "enabled": True, "body": ""},
        {"id": "mediaplan", "title": "Recommended Media Plan", "kind": "mediaplan", "enabled": True, "body": ""},
        {"id": "packages", "title": "Investment Options", "kind": "packages", "enabled": True, "body": ""},
        {"id": "kpis", "title": "How We Measure Success", "kind": "kpis", "enabled": True, "body":
            "Reporting is provided monthly with the KPIs below, reviewed together to optimize the plan."},
        {"id": "landing", "title": "Landing Page Recommendation", "kind": "text", "enabled": True,
         "body": state.get("landing") or ""},
        {"id": "next", "title": "Next Steps", "kind": "text", "enabled": True,
         "body": "1. Approve the recommended package.\n2. Smart 1 finalizes the insertion order and "
                 "creative plan.\n3. Campaign launches after tracking is verified."},
    ]
    return state


# =====================================================================
# AI routes (optional — need OPENAI_API_KEY; same pattern as the IO app)
# =====================================================================
def _openai_response(prompt, max_output_tokens=6000):
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OpenAI is not configured. Add OPENAI_API_KEY.")
    payload = {"model": os.getenv("OPENAI_MODEL", "gpt-5-mini"), "input": prompt,
               "tools": [{"type": "web_search"}], "max_output_tokens": max_output_tokens}
    r = requests.post("https://api.openai.com/v1/responses",
                      headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                      json=payload, timeout=120)
    r.raise_for_status()
    parts = []
    for item in r.json().get("output") or []:
        for content in item.get("content") or []:
            if content.get("type") in ("output_text", "text") and content.get("text"):
                parts.append(content["text"])
    return "\n".join(parts).strip()


def _json_from_ai(text):
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.I | re.S)
    return json.loads(cleaned)


@app.post("/api/ai/draft-proposal")
def ai_draft_proposal():
    """One sentence -> a pre-filled proposal draft the salesperson reviews."""
    body = request.get_json(force=True) or {}
    ask = str(body.get("prompt") or "").strip()
    if not ask:
        return jsonify({"ok": False, "error": "A prompt is required"}), 400
    catalog = body.get("categories") or []
    prompt = (
        "You are a senior media planner at Smart 1 Marketing. From the request below, draft a campaign "
        "proposal setup. Return STRICT JSON only with keys: client (string), url (string, may be empty), "
        "industry (one of: Healthcare / Dental, Home Services, Legal, Automotive, Real Estate, Retail / "
        "E-comm, Restaurant / Hospitality, Financial Services, Fitness / Wellness, Education, B2B / "
        "Professional, Other), objectives (array from: Brand Awareness, Lead Generation, Website Traffic, "
        "Store Visits, Phone Calls, Conversions, Recruitment, Event Promotion), geoType (one of: City/ZIP "
        "+ Radius, DMA, Statewide, National, Other), geo (string), radius (number, only for radius type), "
        "budget (number, monthly dollars), months (number), categories (array of up to 4 product "
        "categories chosen ONLY from this list: " + ", ".join(catalog) + "), rationale (short string).\n\n"
        "Request: " + ask)
    try:
        result = _json_from_ai(_openai_response(prompt, 3000))
        return jsonify({"ok": True, "draft": result})
    except Exception as exc:
        logger.exception("AI draft failed")
        return jsonify({"ok": False, "error": "AI draft failed", "detail": str(exc)}), 502


@app.post("/api/ai/rewrite")
def ai_rewrite():
    """Rewrite a proposal section's copy per an instruction."""
    body = request.get_json(force=True) or {}
    text = str(body.get("text") or "").strip()
    instruction = str(body.get("instruction") or "make it clearer and more client-friendly").strip()
    if not text:
        return jsonify({"ok": False, "error": "Text is required"}), 400
    prompt = ("Rewrite the following marketing-proposal section for Smart 1 Marketing. Instruction: "
              + instruction + ". Keep it factual, professional, and client-facing. Return only the "
              "rewritten text, no preamble.\n\n" + text)
    try:
        return jsonify({"ok": True, "text": _openai_response(prompt, 2500)})
    except Exception as exc:
        return jsonify({"ok": False, "error": "AI rewrite failed", "detail": str(exc)}), 502


# =====================================================================
# TOOLS — SEO Image Optimizer
# ---------------------------------------------------------------------
# Bulk resize + convert to WebP, AI-generated SEO filenames and alt text,
# optional hosting in Cloudinary with the project metadata attached.
#
# Stateless by design: the browser holds the original files and posts them
# once to /analyze (to get names) and again to /finalize or /zip (to save).
# Nothing is cached server-side, so it works across gunicorn workers.
# =====================================================================
MAX_IMAGES_PER_BATCH = int(os.getenv("MAX_IMAGES_PER_BATCH", "10"))
ALLOWED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp", "image/gif", "image/bmp", "image/tiff"}
RESIZE_PRESETS = {
    "hero": 2400,
    "full": 1920,
    "content": 1200,
    "thumb": 800,
    "original": 0,
}


def _slugify(value, fallback="image"):
    text = re.sub(r"[^a-zA-Z0-9]+", "-", str(value or "")).strip("-").lower()
    text = re.sub(r"-{2,}", "-", text)
    return (text or fallback)[:80]


def _fallback_seo(original_name, meta, index):
    """Deterministic naming used when OpenAI is unavailable — never blocks the tool."""
    stem = Path(original_name or "").stem
    generic = bool(re.fullmatch(r"(img|image|photo|dsc|screenshot|untitled)[-_ ]?\d*", stem.strip(), re.I))
    parts = [meta.get("companyName"), meta.get("pageName") or meta.get("projectName")]
    if not generic and stem:
        parts.append(stem)
    slug = _slugify("-".join([p for p in parts if p]), "image")
    if generic or not stem:
        slug = f"{slug}-{index + 1}"
    subject = meta.get("pageName") or meta.get("projectName") or "services"
    alt = f"{meta.get('companyName') or 'Business'} — {subject}".strip()
    return slug, alt[:125]


def _openai_vision_seo(image_bytes, mime_type, meta):
    """Ask OpenAI for an SEO filename + alt text. Raises on any failure."""
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OpenAI is not configured")
    data_url = f"data:{mime_type};base64,{base64.b64encode(image_bytes).decode('ascii')}"
    context_bits = ", ".join(filter(None, [
        f"Company: {meta.get('companyName')}" if meta.get("companyName") else "",
        f"Website: {meta.get('webUrl')}" if meta.get("webUrl") else "",
        f"Project: {meta.get('projectName')}" if meta.get("projectName") else "",
        f"Page: {meta.get('pageName')}" if meta.get("pageName") else "",
    ]))
    payload = {
        "model": os.getenv("OPENAI_VISION_MODEL", "gpt-4o"),
        "messages": [
            {"role": "system", "content":
                "You are an expert SEO copywriter. Analyze images and return JSON with exactly this shape: "
                '{"seoFilename": "descriptive-keyword-rich-filename-lowercase-hyphenated", '
                '"altText": "concise descriptive alt text under 125 characters"}. '
                "The filename must be lowercase, hyphenated, no file extension, no stop-word padding. "
                "The alt text describes what is actually visible; never start it with \"Image of\" or "
                "\"Picture of\", and never stuff keywords."},
            {"role": "user", "content": [
                {"type": "text", "text":
                    "Generate an SEO filename (no extension) and alt text for this image. "
                    f"Business context — {context_bits or 'none provided'}. "
                    "Work the business and page context into the filename naturally where it fits."},
                {"type": "image_url", "image_url": {"url": data_url}},
            ]},
        ],
        "response_format": {"type": "json_object"},
        "max_tokens": 300,
    }
    response = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload, timeout=90,
    )
    response.raise_for_status()
    data = json.loads(response.json()["choices"][0]["message"]["content"])
    slug = _slugify(data.get("seoFilename"), "image")
    alt = str(data.get("altText") or "").strip()[:125]
    if not alt:
        raise RuntimeError("OpenAI returned no alt text")
    return slug, alt


def _read_meta():
    """Project metadata arrives as a JSON string field alongside the files."""
    raw = request.form.get("meta") or "{}"
    try:
        meta = json.loads(raw)
    except Exception:
        meta = {}
    return meta if isinstance(meta, dict) else {}


def _collect_files():
    files = request.files.getlist("images")
    files = [f for f in files if f and f.filename]
    if not files:
        return None, (jsonify({"ok": False, "error": "Select at least one image."}), 400)
    if len(files) > MAX_IMAGES_PER_BATCH:
        return None, (jsonify({"ok": False,
                               "error": f"Up to {MAX_IMAGES_PER_BATCH} images per batch."}), 400)
    for f in files:
        if f.mimetype and f.mimetype not in ALLOWED_IMAGE_TYPES:
            return None, (jsonify({"ok": False,
                                   "error": f"{f.filename} is not a supported image type."}), 400)
    return files, None


def _optimize(image_bytes, max_width, quality):
    """Resize (never upscale) and convert to WebP. Returns (bytes, w, h)."""
    with Image.open(BytesIO(image_bytes)) as img:
        img = ImageOps.exif_transpose(img)          # honor camera rotation
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        if max_width and img.width > max_width:
            ratio = max_width / float(img.width)
            img = img.resize((max_width, max(1, int(img.height * ratio))), Image.LANCZOS)
        out = BytesIO()
        img.save(out, format="WEBP", quality=quality, method=6)
        return out.getvalue(), img.width, img.height


def _quality():
    try:
        return max(40, min(100, int(request.form.get("quality") or 82)))
    except (TypeError, ValueError):
        return 82


def _max_width():
    preset = (request.form.get("preset") or "full").lower()
    if preset == "custom":
        try:
            return max(0, min(6000, int(request.form.get("maxWidth") or 0)))
        except (TypeError, ValueError):
            return 1920
    return RESIZE_PRESETS.get(preset, 1920)


@app.post("/api/tools/image/analyze")
def tools_image_analyze():
    """Step 1: return a suggested SEO filename + alt text for each image."""
    files, error = _collect_files()
    if error:
        return error
    meta = _read_meta()
    ai_key = bool(os.getenv("OPENAI_API_KEY", "").strip())
    results, ai_used, warnings = [], False, []

    for index, storage in enumerate(files):
        raw = storage.read()
        try:
            with Image.open(BytesIO(raw)) as probe:
                width, height = probe.size
        except Exception:
            return jsonify({"ok": False,
                            "error": f"{storage.filename} could not be read as an image."}), 400

        slug, alt, source = None, None, "pattern"
        if ai_key:
            try:
                slug, alt = _openai_vision_seo(raw, storage.mimetype or "image/jpeg", meta)
                source, ai_used = "ai", True
            except Exception as exc:
                logger.warning("Vision naming failed for %s: %s", storage.filename, exc)
                warnings.append(f"AI naming unavailable for {storage.filename} — used a pattern name.")
        if not slug:
            slug, alt = _fallback_seo(storage.filename, meta, index)

        results.append({
            "index": index,
            "originalName": storage.filename,
            "seoFilename": slug,
            "altText": alt,
            "width": width,
            "height": height,
            "sizeKb": round(len(raw) / 1024),
            "source": source,
        })

    return jsonify({
        "ok": True, "files": results, "ai_used": ai_used,
        "ai_available": ai_key,
        "cloudinary_available": _cloudinary_ready(),
        "warnings": warnings,
    })


def _process_approved(files):
    """Shared optimize step for /finalize and /zip. Returns list of dicts."""
    max_width, quality = _max_width(), _quality()
    try:
        approved = json.loads(request.form.get("approved") or "[]")
    except Exception:
        approved = []
    processed = []
    for index, storage in enumerate(files):
        raw = storage.read()
        entry = approved[index] if index < len(approved) and isinstance(approved[index], dict) else {}
        slug = _slugify(entry.get("seoFilename"), f"image-{index + 1}")
        alt = str(entry.get("altText") or "").strip()[:125]
        optimized, width, height = _optimize(raw, max_width, quality)
        processed.append({
            "slug": slug, "altText": alt, "bytes": optimized,
            "width": width, "height": height,
            "originalKb": round(len(raw) / 1024),
            "optimizedKb": round(len(optimized) / 1024),
            "originalName": storage.filename,
        })
    return processed


@app.post("/api/tools/image/zip")
def tools_image_zip():
    """Step 2a: download the optimized WebP files (works without Cloudinary)."""
    files, error = _collect_files()
    if error:
        return error
    meta = _read_meta()
    processed = _process_approved(files)

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        manifest = ["filename,alt_text,width,height,original_kb,optimized_kb"]
        seen = {}
        for item in processed:
            name = item["slug"]
            seen[name] = seen.get(name, 0) + 1
            if seen[name] > 1:
                name = f"{name}-{seen[name]}"
            zf.writestr(f"{name}.webp", item["bytes"])
            manifest.append('{},"{}",{},{},{},{}'.format(
                f"{name}.webp", item["altText"].replace('"', "'"),
                item["width"], item["height"], item["originalKb"], item["optimizedKb"]))
        zf.writestr("alt-text.csv", "\n".join(manifest))
    buf.seek(0)
    project = _slugify(meta.get("projectName"), "images")
    return send_file(buf, mimetype="application/zip", as_attachment=True,
                     download_name=f"S1M Optimized Images - {project}.zip")


@app.post("/api/tools/image/finalize")
def tools_image_finalize():
    """Step 2b: optimize and host in Cloudinary with the project metadata attached."""
    if not _cloudinary_ready():
        return jsonify({"ok": False, "error":
                        "Cloudinary is not configured on this server. Add CLOUDINARY_URL in Render, "
                        "or use Download ZIP instead."}), 503
    files, error = _collect_files()
    if error:
        return error
    meta = _read_meta()
    if not str(meta.get("companyName") or "").strip() or not str(meta.get("projectName") or "").strip():
        return jsonify({"ok": False, "error": "Company name and project name are required."}), 400

    processed = _process_approved(files)
    folder = f"smart1_images/{_slugify(meta.get('companyName'), 'client')}/{_slugify(meta.get('projectName'), 'project')}"
    saved = []
    for item in processed:
        try:
            result = cloudinary.uploader.upload(
                BytesIO(item["bytes"]),
                resource_type="image", folder=folder,
                public_id=item["slug"], overwrite=False, unique_filename=True,
                context={
                    "company": str(meta.get("companyName") or ""),
                    "url": str(meta.get("webUrl") or ""),
                    "project": str(meta.get("projectName") or ""),
                    "page": str(meta.get("pageName") or "N/A"),
                    "alt": item["altText"],
                },
                tags=["smart1_seo_images", _slugify(meta.get("companyName"), "client")],
            )
        except Exception as exc:
            logger.exception("Cloudinary upload failed")
            return jsonify({"ok": False, "error": "Cloudinary upload failed", "detail": str(exc)}), 502
        saved.append({
            "url": result.get("secure_url"),
            "public_id": result.get("public_id"),
            "filename": f"{item['slug']}.webp",
            "altText": item["altText"],
            "width": item["width"], "height": item["height"],
            "originalKb": item["originalKb"], "optimizedKb": item["optimizedKb"],
            "originalName": item["originalName"],
        })

    total_before = sum(s["originalKb"] for s in saved)
    total_after = sum(s["optimizedKb"] for s in saved)
    return jsonify({"ok": True, "savedFiles": saved, "folder": folder,
                    "savedKb": max(0, total_before - total_after),
                    "percentSaved": round(100 * (total_before - total_after) / total_before) if total_before else 0})


@app.get("/api/tools/image/gallery")
def tools_image_gallery():
    """Archive of everything this tool has hosted, newest first."""
    if not _cloudinary_ready():
        return jsonify({"ok": True, "gallery": [], "cloudinary_enabled": False})
    try:
        response = cloudinary.api.resources(
            type="upload", prefix="smart1_images/", context=True,
            max_results=min(int(request.args.get("limit") or 100), 200),
            direction="desc",
        )
    except Exception as exc:
        logger.warning("Cloudinary gallery failed: %s", exc)
        return jsonify({"ok": False, "error": "Could not read the Cloudinary gallery",
                        "detail": str(exc)}), 502
    gallery = []
    for item in response.get("resources", []):
        context = (item.get("context") or {}).get("custom") or item.get("context") or {}
        gallery.append({
            "id": item.get("public_id"),
            "url": item.get("secure_url"),
            "company": context.get("company") or "—",
            "urlRef": context.get("url") or "",
            "project": context.get("project") or "—",
            "page": context.get("page") or "—",
            "alt": context.get("alt") or "",
            "filename": (item.get("public_id") or "").split("/")[-1] + "." + (item.get("format") or "webp"),
            "sizeKb": round((item.get("bytes") or 0) / 1024),
            "width": item.get("width"), "height": item.get("height"),
            "created_at": item.get("created_at"),
        })
    return jsonify({"ok": True, "gallery": gallery, "cloudinary_enabled": True})


@app.errorhandler(413)
def handle_too_large(exc):
    limit = app.config.get("MAX_CONTENT_LENGTH", 0) // (1024 * 1024)
    return jsonify({"ok": False,
                    "error": f"That upload is too large. Keep each batch under {limit} MB."}), 413


@app.errorhandler(Exception)
def handle_unexpected_error(exc):
    logger.exception("Unhandled application error")
    return jsonify({"error": "Internal server error", "type": type(exc).__name__, "message": str(exc)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "8001")), debug=False)
